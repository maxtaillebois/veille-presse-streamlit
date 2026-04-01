#!/usr/bin/env python3
"""
Application de Veille Presse v5.0 — Cloud Edition + USH Panorama
=================================================================
Deployable sur Streamlit Community Cloud.
Recupere les donnees (JSON + PDF) depuis Google Drive.

Onglet 1 — Veille interne : articles PDF analyses par IA (workflow N8N 2)
Onglet 2 — USH Panorama   : articles du panorama Cision (workflow N8N 4)

Les deux sources sont selectionnables et compilables ensemble
pour envoi a Stephanie via le workflow N8N 3.

Lancer en local : streamlit run veille_presse_app.py
"""

import streamlit as st
from pypdf import PdfReader, PdfWriter
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re
import io
import json
import zipfile
import requests
from datetime import datetime


# ==================== CONFIGURATION ====================

# ID du dossier Google Drive partage (contient les PDF + veille_data.json + ush_panorama_data.json)
GDRIVE_FOLDER_ID = "15MwTpntChfAgDQYiDSvc5cJ0KYmuW9_p"

# URL du webhook N8N pour l'envoi du mail a Stephanie
N8N_WEBHOOK_URL = "https://maximetaillebois.app.n8n.cloud/webhook/veille-presse-envoi"

# Mots-cles de reference (pour le badge couleur — veille interne)
KEYWORDS = ["Yannick Borde", "Procivis", "Immo de France", "Maisons d'en France"]


def get_google_api_key():
    """Recupere la cle API Google depuis les secrets Streamlit ou l'environnement."""
    try:
        return st.secrets["GOOGLE_API_KEY"]
    except Exception:
        return os.environ.get("GOOGLE_API_KEY", "")


# ==================== FONCTIONS GOOGLE DRIVE ====================

@st.cache_data(ttl=300)  # Cache 5 minutes
def list_drive_files(folder_id: str, api_key: str) -> list:
    """Liste les fichiers dans un dossier Google Drive partage."""
    url = "https://www.googleapis.com/drive/v3/files"
    params = {
        "q": f"'{folder_id}' in parents and trashed = false",
        "fields": "files(id,name,mimeType,modifiedTime,size)",
        "orderBy": "modifiedTime desc",
        "pageSize": 100,
        "key": api_key
    }
    try:
        resp = requests.get(url, params=params, timeout=15)
        resp.raise_for_status()
        return resp.json().get("files", [])
    except Exception as e:
        st.error(f"Erreur lors de la lecture du Google Drive : {e}")
        return []


@st.cache_data(ttl=300)
def download_drive_file(file_id: str, api_key: str) -> bytes:
    """Telecharge un fichier depuis Google Drive."""
    url = f"https://www.googleapis.com/drive/v3/files/{file_id}"
    params = {"alt": "media", "key": api_key}
    try:
        resp = requests.get(url, params=params, timeout=60)
        resp.raise_for_status()
        return resp.content
    except Exception as e:
        st.error(f"Erreur de telechargement : {e}")
        return b""


def find_file_in_drive(files: list, filename: str) -> dict:
    """Trouve un fichier par son nom dans la liste Drive."""
    for f in files:
        if f["name"] == filename:
            return f
    return {}


def load_json_from_drive(files: list, api_key: str, filename: str) -> dict:
    """Charge un fichier JSON depuis Google Drive."""
    json_file = find_file_in_drive(files, filename)
    if not json_file:
        return None
    content = download_drive_file(json_file["id"], api_key)
    if not content:
        return None
    try:
        return json.loads(content.decode("utf-8"))
    except json.JSONDecodeError as e:
        st.error(f"Erreur de lecture du fichier {filename} : {e}")
        return None


# ==================== CONFIGURATION PAGE ====================

st.set_page_config(
    page_title="Veille Presse — Procivis",
    page_icon="📰",
    layout="wide"
)

st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1E3A5F;
        margin-bottom: 0.5rem;
    }
    .sub-header {
        font-size: 1.1rem;
        color: #666;
        margin-bottom: 2rem;
    }
    .keyword-highlight {
        background-color: #FFEB3B;
        padding: 2px 4px;
        border-radius: 3px;
        font-weight: bold;
    }
    .media-badge {
        background-color: #1976D2;
        color: white;
        padding: 3px 10px;
        border-radius: 12px;
        font-size: 0.85rem;
        font-weight: bold;
        display: inline-block;
    }
    .date-badge {
        background-color: #43A047;
        color: white;
        padding: 3px 10px;
        border-radius: 12px;
        font-size: 0.85rem;
        display: inline-block;
    }
    .keyword-badge {
        background-color: #FF9800;
        color: white;
        padding: 3px 10px;
        border-radius: 12px;
        font-size: 0.85rem;
        display: inline-block;
    }
    .theme-badge {
        background-color: #E43535;
        color: white;
        padding: 3px 10px;
        border-radius: 12px;
        font-size: 0.85rem;
        font-weight: bold;
        display: inline-block;
    }
    .type-badge {
        background-color: #7B1FA2;
        color: white;
        padding: 3px 10px;
        border-radius: 12px;
        font-size: 0.85rem;
        display: inline-block;
    }
    .author-badge {
        background-color: #546E7A;
        color: white;
        padding: 3px 10px;
        border-radius: 12px;
        font-size: 0.85rem;
        display: inline-block;
    }
    .context-text {
        font-style: italic;
        color: #424242;
        background-color: #F5F5F5;
        padding: 10px;
        border-radius: 5px;
        margin: 8px 0;
        border-left: 3px solid #1976D2;
    }
    .summary-text {
        color: #333;
        background-color: #E8F5E9;
        padding: 10px;
        border-radius: 5px;
        margin: 8px 0;
        border-left: 3px solid #43A047;
    }
    .ush-link {
        color: #1976D2;
        text-decoration: none;
        font-weight: bold;
    }
    .stButton > button { width: 100%; }
</style>
""", unsafe_allow_html=True)


# ==================== FONCTIONS UTILITAIRES ====================

def highlight_keywords_html(text: str, keywords: list) -> str:
    """Surligne les mots-cles dans le texte (HTML)."""
    result = text
    for keyword in keywords:
        pattern = re.compile(re.escape(keyword), re.IGNORECASE)
        result = pattern.sub(
            f'<span class="keyword-highlight">{keyword}</span>',
            result
        )
    return result


def parse_date_for_sorting(date_str: str) -> datetime:
    """Parse une date JJ/MM/AAAA ou autre format pour le tri."""
    if not date_str:
        return datetime.min
    # Nettoyer la date
    cleaned = date_str.strip()
    # Formats courants
    for fmt in ["%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y", "%d/%m/%y",
                "%d/%m/%y à %H:%M", "%d/%m/%y \u00e0 %H:%M"]:
        try:
            return datetime.strptime(cleaned, fmt)
        except ValueError:
            continue
    # Format textuel francais : "1er avril 2026", "31 mars 2026"
    mois_fr = {
        'janvier': 1, 'février': 2, 'fevrier': 2, 'mars': 3,
        'avril': 4, 'mai': 5, 'juin': 6, 'juillet': 7,
        'août': 8, 'aout': 8, 'septembre': 9, 'octobre': 10,
        'novembre': 11, 'décembre': 12, 'decembre': 12
    }
    try:
        parts = cleaned.lower().replace('1er', '1').split()
        if len(parts) >= 3:
            jour = int(parts[0])
            mois = mois_fr.get(parts[1], 0)
            annee = int(parts[2])
            if mois > 0:
                return datetime(annee, mois, jour)
    except (ValueError, IndexError):
        pass
    return datetime.min


# ==================== FONCTIONS DE COMPILATION ====================

def create_compilation_pdf_bytes(pdf_contents: list) -> bytes:
    """Compile plusieurs PDF (en bytes) en un seul fichier."""
    writer = PdfWriter()
    for pdf_bytes, filename in pdf_contents:
        try:
            reader = PdfReader(io.BytesIO(pdf_bytes))
            for page in reader.pages:
                writer.add_page(page)
        except Exception as e:
            st.warning(f"Impossible d'ajouter {filename} : {e}")

    output = io.BytesIO()
    writer.write(output)
    output.seek(0)
    return output.getvalue()


def create_recap_docx_bytes(veille_articles: list, ush_articles: list = None) -> bytes:
    """Cree le Word recapitulatif en memoire, incluant les deux sources."""
    doc = DocxDocument()

    title = doc.add_heading('RECAPITULATIF VEILLE PRESSE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(
        f"Genere le {datetime.now().strftime('%d/%m/%Y a %H:%M')}"
    )
    info_run.font.size = Pt(10)
    info_run.font.italic = True

    total = len(veille_articles) + (len(ush_articles) if ush_articles else 0)
    count_para = doc.add_paragraph()
    count_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    count_run = count_para.add_run(f"{total} article(s)")
    count_run.font.size = Pt(10)

    doc.add_paragraph("-" * 50)

    # Section veille interne
    if veille_articles:
        doc.add_heading('VEILLE INTERNE (Sophie)', level=1)
        sorted_articles = sorted(
            veille_articles,
            key=lambda x: parse_date_for_sorting(x.get('date', '')),
            reverse=True
        )
        for article in sorted_articles:
            media = article.get('media', 'Media inconnu')
            title_text = article.get('title', 'Titre inconnu')
            date = article.get('date', 'Date inconnue')
            para = doc.add_paragraph()
            media_run = para.add_run(media)
            media_run.bold = True
            para.add_run(" | ")
            title_run = para.add_run(f'"{title_text}"')
            title_run.italic = True
            para.add_run(" | ")
            para.add_run(date)

    # Section USH Panorama
    if ush_articles:
        doc.add_paragraph("-" * 50)
        doc.add_heading('PANORAMA USH (Cision)', level=1)
        sorted_ush = sorted(
            ush_articles,
            key=lambda x: parse_date_for_sorting(x.get('date', '')),
            reverse=True
        )
        for article in sorted_ush:
            media = article.get('media', 'Media inconnu')
            title_text = article.get('title', 'Titre inconnu')
            date = article.get('date', 'Date inconnue')
            media_type = article.get('media_type', '')
            para = doc.add_paragraph()
            media_run = para.add_run(media)
            media_run.bold = True
            if media_type:
                para.add_run(f" [{media_type}]")
            para.add_run(" | ")
            title_run = para.add_run(f'"{title_text}"')
            title_run.italic = True
            para.add_run(" | ")
            para.add_run(date)
            author = article.get('author', '')
            if author:
                para.add_run(f" | {author}")

    doc.add_paragraph("-" * 50)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


# ==================== FONCTION WEBHOOK ====================

def send_to_stephanie(pdf_bytes: bytes, docx_bytes: bytes,
                      veille_articles: list, ush_articles: list,
                      week_start: str, week_end: str) -> bool:
    """Declenche le webhook N8N pour envoyer le mail a Stephanie."""
    if not N8N_WEBHOOK_URL:
        st.error("L'URL du webhook N8N n'est pas configuree.")
        return False

    try:
        date_suffix = datetime.now().strftime('%Y%m%d')
        files = {}
        if pdf_bytes:
            files['pdf'] = (
                f"compilation_veille_{date_suffix}.pdf",
                io.BytesIO(pdf_bytes),
                'application/pdf'
            )
        if docx_bytes:
            files['docx'] = (
                f"compilation_veille_{date_suffix}_recap.docx",
                io.BytesIO(docx_bytes),
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

        # Combiner les articles des deux sources pour le mail
        all_articles = []
        for a in veille_articles:
            all_articles.append({
                "media": a.get("media", ""),
                "title": a.get("title", ""),
                "date": a.get("date", ""),
                "source": "Veille interne"
            })
        for a in (ush_articles or []):
            all_articles.append({
                "media": a.get("media", ""),
                "title": a.get("title", ""),
                "date": a.get("date", ""),
                "source": "USH Panorama"
            })

        data = {
            'week_start': week_start,
            'week_end': week_end,
            'article_count': str(len(all_articles)),
            'generated_at': datetime.now().isoformat(),
            'selected_articles': json.dumps(all_articles)
        }

        response = requests.post(N8N_WEBHOOK_URL, data=data, files=files, timeout=30)
        return response.status_code == 200

    except Exception as e:
        st.error(f"Erreur lors de l'envoi : {e}")
        return False


# ==================== ONGLET 1 : VEILLE INTERNE ====================

def render_veille_interne(drive_files, api_key, pdf_drive_index):
    """Affiche l'onglet Veille Interne (articles PDF / Sophie)."""

    data = load_json_from_drive(drive_files, api_key, "veille_data.json")

    if data is None:
        st.info(
            "En attente du fichier `veille_data.json` dans Google Drive.\n\n"
            "Le workflow N8N s'execute chaque vendredi a 9h00."
        )
        return [], '', ''

    articles = data.get('articles', [])
    week_start = data.get('week_start', '')
    week_end = data.get('week_end', '')
    generated_at = data.get('generated_at', '')

    if not articles:
        st.warning("Le fichier JSON ne contient aucun article.")
        return [], week_start, week_end

    # En-tete
    if week_start and week_end:
        st.markdown(
            f'<p class="sub-header">Semaine du {week_start} au {week_end} '
            f'| {len(articles)} article(s) analyse(s) '
            f'| Donnees du {generated_at[:10] if generated_at else "?"}</p>',
            unsafe_allow_html=True
        )
    else:
        st.markdown(
            f'<p class="sub-header">{len(articles)} article(s) analyse(s) '
            f'| Donnees du {generated_at[:10] if generated_at else "?"}</p>',
            unsafe_allow_html=True
        )

    # Statistiques
    col_stat1, col_stat2, col_stat3, col_stat4 = st.columns(4)
    with col_stat1:
        st.metric("Articles", len(articles))
    with col_stat2:
        kw_counts = {}
        for a in articles:
            for kw in a.get('keywords_found', []):
                kw_counts[kw] = kw_counts.get(kw, 0) + 1
        top_kw = max(kw_counts, key=kw_counts.get) if kw_counts else "—"
        st.metric("Mot-cle principal", top_kw)
    with col_stat3:
        medias = set(a.get('media', '') for a in articles)
        st.metric("Medias distincts", len(medias))
    with col_stat4:
        selected = sum(
            1 for i, a in enumerate(articles)
            if a.get('id', f'veille_{i}') in st.session_state.selected_veille
        )
        st.metric("Selectionnes", f"{selected}/{len(articles)}")

    st.divider()

    # Boutons de selection rapide
    col_sel1, col_sel2, col_sel3 = st.columns([1, 1, 3])
    with col_sel1:
        if st.button("Tout selectionner", key="sel_all_veille", use_container_width=True):
            for i, a in enumerate(articles):
                st.session_state.selected_veille.add(a.get('id', f'veille_{i}'))
            st.rerun()
    with col_sel2:
        if st.button("Tout deselectionner", key="desel_all_veille", use_container_width=True):
            st.session_state.selected_veille = set()
            st.rerun()

    st.divider()

    # Affichage des articles
    for i, article in enumerate(articles):
        article_id = article.get('id', f'veille_{i}')
        media = article.get('media', 'Media inconnu')
        title = article.get('title', 'Titre inconnu')
        date = article.get('date', 'Date inconnue')
        keywords_found = article.get('keywords_found', [])
        context = article.get('context', '')
        summary = article.get('summary', '')
        pdf_filename = article.get('pdf_filename', article.get('fileName', ''))

        with st.container():
            col_main, col_select = st.columns([5, 1])

            with col_main:
                st.markdown(f"### {title}")

                badges_html = f'<span class="media-badge">{media}</span> '
                badges_html += f'<span class="date-badge">{date}</span> '
                for kw in keywords_found:
                    badges_html += f'<span class="keyword-badge">{kw}</span> '
                st.markdown(badges_html, unsafe_allow_html=True)

                if pdf_filename:
                    has_pdf = pdf_filename in pdf_drive_index
                    icon = "📎" if has_pdf else "⚠️"
                    st.caption(f"{icon} {pdf_filename}")

                if summary:
                    st.markdown(
                        f'<div class="summary-text"><b>Resume :</b> {summary}</div>',
                        unsafe_allow_html=True
                    )

                if context:
                    highlighted = highlight_keywords_html(context, keywords_found)
                    with st.expander("Voir le contexte de citation", expanded=False):
                        st.markdown(
                            f'<div class="context-text">{highlighted}</div>',
                            unsafe_allow_html=True
                        )

            with col_select:
                is_selected = st.checkbox(
                    "Garder",
                    key=f"select_veille_{article_id}",
                    value=article_id in st.session_state.selected_veille
                )
                if is_selected:
                    st.session_state.selected_veille.add(article_id)
                else:
                    st.session_state.selected_veille.discard(article_id)

            st.divider()

    # Retourner les articles, week_start, week_end pour la compilation
    return articles, week_start, week_end


# ==================== ONGLET 2 : USH PANORAMA ====================

def render_ush_panorama(drive_files, api_key):
    """Affiche l'onglet USH Panorama de Presse (Cision)."""

    data = load_json_from_drive(drive_files, api_key, "ush_panorama_data.json")

    if data is None:
        st.info(
            "En attente du fichier `ush_panorama_data.json` dans Google Drive.\n\n"
            "Le workflow N8N 4 s'execute chaque jour a la reception "
            "du mail Cision."
        )
        return []

    articles = data.get('articles', [])
    panorama_date = data.get('panorama_date', '')
    generated_at = data.get('generated_at', '')
    themes = data.get('themes', [])

    if not articles:
        st.warning("Le panorama USH ne contient aucun article.")
        return []

    # En-tete
    st.markdown(
        f'<p class="sub-header">Panorama du {panorama_date} '
        f'| {len(articles)} article(s) '
        f'| {len(themes)} theme(s) '
        f'| Donnees du {generated_at[:10] if generated_at else "?"}</p>',
        unsafe_allow_html=True
    )

    # Statistiques
    col_stat1, col_stat2, col_stat3, col_stat4 = st.columns(4)
    with col_stat1:
        st.metric("Articles", len(articles))
    with col_stat2:
        types_count = {}
        for a in articles:
            t = a.get('media_type', 'Autre')
            types_count[t] = types_count.get(t, 0) + 1
        type_summary = ", ".join(f"{v} {k}" for k, v in types_count.items())
        st.metric("Types", type_summary if type_summary else "—")
    with col_stat3:
        medias = set(a.get('media', '') for a in articles)
        st.metric("Medias distincts", len(medias))
    with col_stat4:
        selected = sum(
            1 for a in articles
            if a.get('id', '') in st.session_state.selected_ush
        )
        st.metric("Selectionnes", f"{selected}/{len(articles)}")

    st.divider()

    # Boutons de selection rapide
    col_sel1, col_sel2, col_sel3 = st.columns([1, 1, 3])
    with col_sel1:
        if st.button("Tout selectionner", key="sel_all_ush", use_container_width=True):
            for a in articles:
                st.session_state.selected_ush.add(a.get('id', ''))
            st.rerun()
    with col_sel2:
        if st.button("Tout deselectionner", key="desel_all_ush", use_container_width=True):
            st.session_state.selected_ush = set()
            st.rerun()

    # Filtre par theme
    with col_sel3:
        theme_filter = st.selectbox(
            "Filtrer par theme",
            options=["Tous"] + themes,
            key="ush_theme_filter"
        )

    st.divider()

    # Affichage des articles
    displayed_articles = articles
    if theme_filter != "Tous":
        displayed_articles = [a for a in articles if a.get('theme') == theme_filter]

    for article in displayed_articles:
        article_id = article.get('id', '')
        media = article.get('media', 'Media inconnu')
        title = article.get('title', 'Titre inconnu')
        date = article.get('date', 'Date inconnue')
        author = article.get('author', '')
        media_type = article.get('media_type', '')
        theme = article.get('theme', '')
        link = article.get('link', '')

        with st.container():
            col_main, col_select = st.columns([5, 1])

            with col_main:
                # Titre avec lien vers LuQi si disponible
                if link:
                    st.markdown(f"### [{title}]({link})")
                else:
                    st.markdown(f"### {title}")

                # Badges
                badges_html = f'<span class="media-badge">{media}</span> '
                badges_html += f'<span class="date-badge">{date}</span> '
                if media_type:
                    badges_html += f'<span class="type-badge">{media_type}</span> '
                if theme:
                    badges_html += f'<span class="theme-badge">{theme}</span> '
                if author:
                    badges_html += f'<span class="author-badge">{author}</span> '
                st.markdown(badges_html, unsafe_allow_html=True)

            with col_select:
                is_selected = st.checkbox(
                    "Garder",
                    key=f"select_ush_{article_id}",
                    value=article_id in st.session_state.selected_ush
                )
                if is_selected:
                    st.session_state.selected_ush.add(article_id)
                else:
                    st.session_state.selected_ush.discard(article_id)

            st.divider()

    return articles


# ==================== INTERFACE PRINCIPALE ====================

def main():
    st.markdown(
        '<h1 class="main-header">📰 Veille Presse</h1>',
        unsafe_allow_html=True
    )

    # Verifier la cle API
    api_key = get_google_api_key()
    if not api_key:
        st.error(
            "**Cle API Google manquante.**\n\n"
            "Configurez `GOOGLE_API_KEY` dans les secrets Streamlit "
            "(Settings > Secrets) ou en variable d'environnement."
        )
        st.code('GOOGLE_API_KEY = "votre-cle-api-google"', language="toml")
        return

    # ===== SIDEBAR =====
    with st.sidebar:
        st.header("Configuration")

        st.caption(f"Dossier Google Drive : `{GDRIVE_FOLDER_ID}`")

        if st.button("Rafraichir les donnees", use_container_width=True):
            st.cache_data.clear()
            st.rerun()

        st.divider()

        st.subheader("Envoi mail")
        webhook_url = st.text_input(
            "URL webhook N8N",
            value=N8N_WEBHOOK_URL,
            help="URL du webhook N8N pour l'envoi automatique",
            type="password"
        )

    # ===== INITIALISATION SESSION STATE =====
    if 'selected_veille' not in st.session_state:
        st.session_state.selected_veille = set()
    if 'selected_ush' not in st.session_state:
        st.session_state.selected_ush = set()
    if 'compilation_done' not in st.session_state:
        st.session_state.compilation_done = False
    if 'compilation_data' not in st.session_state:
        st.session_state.compilation_data = {}

    # ===== CHARGEMENT DES FICHIERS DRIVE =====
    with st.spinner("Chargement depuis Google Drive..."):
        drive_files = list_drive_files(GDRIVE_FOLDER_ID, api_key)

    if not drive_files:
        st.info(
            "Aucun fichier trouve dans le dossier Google Drive.\n\n"
            "Les workflows N8N deposent les fichiers automatiquement."
        )
        return

    # Indexer les PDF du Drive par nom de fichier
    pdf_drive_index = {}
    for f in drive_files:
        if f["name"].lower().endswith(".pdf"):
            pdf_drive_index[f["name"]] = f["id"]

    # ===== ONGLETS =====
    tab_veille, tab_ush, tab_compile = st.tabs([
        "📋 Veille interne",
        "🏛️ USH Panorama",
        "📦 Compilation & Envoi"
    ])

    with tab_veille:
        veille_articles, week_start, week_end = render_veille_interne(
            drive_files, api_key, pdf_drive_index
        )

    with tab_ush:
        ush_articles = render_ush_panorama(drive_files, api_key)

    # ===== ONGLET COMPILATION =====
    with tab_compile:
        st.header("Compilation & Envoi")

        # Compter les selections
        selected_veille_list = []
        for i, a in enumerate(veille_articles):
            aid = a.get('id', f'veille_{i}')
            if aid in st.session_state.selected_veille:
                selected_veille_list.append(a)

        selected_ush_list = []
        for a in ush_articles:
            aid = a.get('id', '')
            if aid in st.session_state.selected_ush:
                selected_ush_list.append(a)

        total_selected = len(selected_veille_list) + len(selected_ush_list)

        # Resume de la selection
        col_r1, col_r2, col_r3 = st.columns(3)
        with col_r1:
            st.metric("Veille interne", f"{len(selected_veille_list)} article(s)")
        with col_r2:
            st.metric("USH Panorama", f"{len(selected_ush_list)} article(s)")
        with col_r3:
            st.metric("Total", f"{total_selected} article(s)")

        st.divider()

        if total_selected > 0:
            # Afficher le recap des articles selectionnes
            if selected_veille_list:
                st.subheader("Veille interne")
                for a in selected_veille_list:
                    st.markdown(
                        f"- **{a.get('media', '?')}** | "
                        f"*\"{a.get('title', '?')}\"* | "
                        f"{a.get('date', '?')}"
                    )

            if selected_ush_list:
                st.subheader("USH Panorama")
                for a in selected_ush_list:
                    st.markdown(
                        f"- **{a.get('media', '?')}** [{a.get('media_type', '')}] | "
                        f"*\"{a.get('title', '?')}\"* | "
                        f"{a.get('date', '?')}"
                    )

            st.divider()

            date_suffix = datetime.now().strftime('%Y%m%d')
            col_compile, col_send = st.columns(2)

            with col_compile:
                if st.button("Generer la compilation", type="primary", use_container_width=True):
                    with st.spinner("Generation en cours..."):
                        # 1. Word recapitulatif (les deux sources)
                        docx_bytes = create_recap_docx_bytes(
                            selected_veille_list, selected_ush_list
                        )

                        # 2. PDF : seulement les articles de la veille interne
                        pdf_contents = []
                        missing_pdfs = []
                        for a in selected_veille_list:
                            pdf_filename = a.get('pdf_filename', a.get('fileName', ''))
                            if pdf_filename and pdf_filename in pdf_drive_index:
                                file_id = pdf_drive_index[pdf_filename]
                                pdf_bytes = download_drive_file(file_id, api_key)
                                if pdf_bytes:
                                    pdf_contents.append((pdf_bytes, pdf_filename))
                                else:
                                    missing_pdfs.append(pdf_filename)
                            else:
                                missing_pdfs.append(pdf_filename or a.get('title', '?'))

                        compiled_pdf_bytes = b""
                        if pdf_contents:
                            compiled_pdf_bytes = create_compilation_pdf_bytes(pdf_contents)

                        if missing_pdfs:
                            st.warning(
                                f"{len(missing_pdfs)} PDF non trouve(s) : "
                                f"{', '.join(missing_pdfs)}"
                            )

                        st.session_state.compilation_done = True
                        st.session_state.compilation_data = {
                            'pdf_bytes': compiled_pdf_bytes,
                            'docx_bytes': docx_bytes,
                        }

                        st.success("Recapitulatif Word genere")
                        if compiled_pdf_bytes:
                            st.success(
                                f"Compilation PDF generee ({len(pdf_contents)} articles)"
                            )
                        if selected_ush_list:
                            st.info(
                                f"{len(selected_ush_list)} article(s) USH inclus "
                                f"dans le recapitulatif Word (pas de PDF pour ces articles)"
                            )
                        st.balloons()

            # Boutons de telechargement
            if st.session_state.compilation_done:
                comp_data = st.session_state.compilation_data

                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
                    if comp_data.get('docx_bytes'):
                        zf.writestr(
                            f"compilation_veille_{date_suffix}_recap.docx",
                            comp_data['docx_bytes']
                        )
                    if comp_data.get('pdf_bytes'):
                        zf.writestr(
                            f"compilation_veille_{date_suffix}.pdf",
                            comp_data['pdf_bytes']
                        )
                zip_buffer.seek(0)

                st.download_button(
                    label="Telecharger tout (PDF + Word)",
                    data=zip_buffer.getvalue(),
                    file_name=f"veille_presse_{date_suffix}.zip",
                    mime="application/zip",
                    use_container_width=True
                )

            # Bouton d'envoi a Stephanie
            with col_send:
                if st.session_state.compilation_done:
                    if st.button(
                        "Envoyer a Stephanie",
                        type="secondary",
                        use_container_width=True
                    ):
                        comp_data = st.session_state.compilation_data
                        effective_webhook = webhook_url or N8N_WEBHOOK_URL

                        if not effective_webhook:
                            st.error("L'URL du webhook N8N n'est pas configuree.")
                        else:
                            with st.spinner("Envoi en cours..."):
                                success = send_to_stephanie(
                                    pdf_bytes=comp_data.get('pdf_bytes', b''),
                                    docx_bytes=comp_data.get('docx_bytes', b''),
                                    veille_articles=selected_veille_list,
                                    ush_articles=selected_ush_list,
                                    week_start=week_start if veille_articles else '',
                                    week_end=week_end if veille_articles else ''
                                )
                                if success:
                                    st.success("Mail envoye a Stephanie !")
                                else:
                                    st.error(
                                        "L'envoi a echoue. Verifiez le webhook N8N."
                                    )
                else:
                    st.button(
                        "Envoyer a Stephanie",
                        type="secondary",
                        use_container_width=True,
                        disabled=True,
                        help="Generez d'abord la compilation"
                    )
        else:
            st.warning(
                "Selectionnez au moins un article (dans l'onglet Veille interne "
                "ou USH Panorama) pour creer une compilation."
            )


if __name__ == "__main__":
    main()
